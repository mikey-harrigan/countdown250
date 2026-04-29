"""Convert the Hosted Tables docs (markdown) to Word .docx files.

Targeted at the markdown patterns actually used in docs/*.md:
  - ATX headings (# / ## / ###)
  - Pipe tables with a separator row
  - Bulleted lists (- or *) and numbered lists (1.)
  - Inline emphasis: **bold**, *italic*, `code`
  - Inline links: [text](url)
  - Fenced code blocks (```)
  - Blockquotes (>)
  - Horizontal rules (---)
  - Plain paragraphs

Run from repo root:
  python scripts/md_to_docx.py
Outputs to: docs/word/*.docx
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT      = Path(r"C:/Users/USER/countdown250")
SRC_DIR   = ROOT / "docs"
OUT_DIR   = ROOT / "docs" / "word"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# Files to convert; first pair is the most important (the manager handoff).
FILES = [
    "hosted-tables-handoff.md",
    "hosted-tables-program.md",
    "website-design-conventions.md",
]

# ---------- Inline-formatting parser (bold, italic, code, links) ----------

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*"          # **bold**
    r"|`[^`]+`"                # `code`
    r"|\*[^*]+\*"              # *italic*
    r"|\[[^\]]+\]\([^)]+\))"   # [text](url)
)
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _set_cell_shading(cell, hex_color: str):
    """Set background fill on a table cell (e.g., header row)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_runs_with_inline_formatting(paragraph, text: str):
    """Append runs to `paragraph` interpreting **, *, `, [text](url)."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("["):
            lm = LINK_RE.match(token)
            if lm:
                # Render as the visible text in a distinct color (no clickable
                # hyperlink: keeping it simple is fine for this handoff doc).
                r = paragraph.add_run(lm.group(1))
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                r.font.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------- Markdown line classifier ----------

H_RE       = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")
HR_RE      = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
ULI_RE     = re.compile(r"^(\s*)[-*]\s+(.+)$")
OLI_RE     = re.compile(r"^(\s*)\d+\.\s+(.+)$")
BQ_RE      = re.compile(r"^\s*>\s?(.*)$")
TABLE_ROW  = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP  = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$")
FENCE_RE   = re.compile(r"^```")


def parse_table_rows(rows: list[str]) -> tuple[list[str], list[list[str]]]:
    """Given pipe-table lines (header, separator, data...), return (header, data-rows)."""
    def split_row(line: str) -> list[str]:
        # Strip leading/trailing pipes, split on |, trim each cell
        s = line.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    header = split_row(rows[0])
    data = [split_row(r) for r in rows[2:]]  # skip the separator at index 1
    return header, data


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Default style: tighter, more readable
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading styles: keep punchy
    for lvl, sz in ((1, 22), (2, 16), (3, 13), (4, 12)):
        s = doc.styles[f"Heading {lvl}"]
        s.font.size = Pt(sz)
        s.font.color.rgb = RGBColor(0x0B, 0x1F, 0x3A)  # navy
        s.font.bold = True

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip blank
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(stripped):
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            p_pr.append(pBdr)
            i += 1
            continue

        # Fenced code block
        if FENCE_RE.match(stripped):
            i += 1
            buf = []
            while i < n and not FENCE_RE.match(lines[i].strip()):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if buf:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run("\n".join(buf))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                # Light grey shading
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                p_pr.append(shd)
            continue

        # Heading
        m_h = H_RE.match(stripped)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {level}"]
            add_runs_with_inline_formatting(p, m_h.group(2))
            i += 1
            continue

        # Pipe table — detect by current line being a row AND next line being separator
        if TABLE_ROW.match(stripped) and i + 1 < n and TABLE_SEP.match(lines[i + 1].strip()):
            row_lines = []
            while i < n and TABLE_ROW.match(lines[i].strip()):
                row_lines.append(lines[i])
                i += 1
            if len(row_lines) >= 2:
                header, data_rows = parse_table_rows(row_lines)
                ncols = len(header)
                # Ensure all rows have ncols (pad short rows)
                norm = lambda row: row + [""] * (ncols - len(row))
                table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                # Header
                for c, h in enumerate(header):
                    cell = table.rows[0].cells[c]
                    cell.text = ""  # clear default
                    p = cell.paragraphs[0]
                    add_runs_with_inline_formatting(p, h)
                    for run in p.runs:
                        run.bold = True
                    _set_cell_shading(cell, "0B1F3A")
                    # White text on navy header
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Body
                for r_idx, row in enumerate(data_rows, start=1):
                    for c_idx, txt in enumerate(norm(row)):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs_with_inline_formatting(p, txt)
                doc.add_paragraph()  # spacer after table
            continue

        # Blockquote
        if BQ_RE.match(stripped):
            buf = []
            while i < n and BQ_RE.match(lines[i].strip()) and lines[i].strip():
                m = BQ_RE.match(lines[i].strip())
                buf.append(m.group(1))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            add_runs_with_inline_formatting(p, " ".join(buf))
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Bulleted list
        m_uli = ULI_RE.match(line)
        if m_uli:
            while i < n and (ULI_RE.match(lines[i]) or (lines[i].strip() == "")):
                if lines[i].strip() == "":
                    break
                m = ULI_RE.match(lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_inline_formatting(p, m.group(2))
                i += 1
            continue

        # Numbered list
        m_oli = OLI_RE.match(line)
        if m_oli:
            while i < n and (OLI_RE.match(lines[i]) or (lines[i].strip() == "")):
                if lines[i].strip() == "":
                    break
                m = OLI_RE.match(lines[i])
                p = doc.add_paragraph(style="List Number")
                add_runs_with_inline_formatting(p, m.group(2))
                i += 1
            continue

        # Plain paragraph (gather consecutive non-blank, non-special lines)
        buf = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            ns = nxt.strip()
            if not ns:
                break
            if (HR_RE.match(ns) or H_RE.match(ns) or FENCE_RE.match(ns)
                or ULI_RE.match(nxt) or OLI_RE.match(nxt) or BQ_RE.match(ns)
                or (TABLE_ROW.match(ns) and i + 1 < n and TABLE_SEP.match(lines[i + 1].strip()))):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_runs_with_inline_formatting(p, " ".join(b.strip() for b in buf))

    doc.save(docx_path)


# ---------- Run ----------

if __name__ == "__main__":
    for fname in FILES:
        src = SRC_DIR / fname
        dst = OUT_DIR / (src.stem + ".docx")
        if not src.exists():
            print(f"MISSING: {src}")
            continue
        convert_md_to_docx(src, dst)
        kb = dst.stat().st_size / 1024
        print(f"  {src.name:42s} -> {dst.relative_to(ROOT)}  ({kb:6.1f} KB)")
    print("\nDone.")
